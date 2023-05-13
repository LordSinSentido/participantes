from docx import Document
from docx2pdf import convert
import docx2txt
import os


class Exporter:
    def __init__(self):
        pass

    def getTicket(self, data, extension="DOCX"):
        try:
            document = Document("drafts/ticket.docx")
            table = document.tables[0]

            # Agrega el nombre completo
            fullName = f"{data[0]} {data[1]} {data[2]}"
            table.cell(0, 0).paragraphs[1].text = fullName

            # Agrega la edad
            table.cell(1, 0).paragraphs[1].text = f"{data[3]}"

            # Agrega el sexo
            table.cell(1, 1).paragraphs[1].text = f"{data[4]}"

            # Agrega el curp
            table.cell(2, 0).paragraphs[1].text = f"{data[7]}"

            # Agrega la dirección
            table.cell(3, 0).paragraphs[1].text = f"{data[6]}"

            # Agrega la categoría
            table.cell(4, 0).paragraphs[1].text = f"{data[8]}"

            # Agrega la escuela
            table.cell(
                5, 0).paragraphs[1].text = f"{data[5]}" if data[5] != "" else "No aplica"

            # Agrega el pago total
            table.cell(6, 0).paragraphs[1].text = f"${data[9]}"

            if extension == "DOCX":
                document.save(f"saves/tickets/{fullName}.docx")

            else:
                document.save(f"temp/{fullName}.docx")

                if extension == "PDF":
                    convert(f"temp/{fullName}.docx",
                            f"saves/tickets/{fullName}.pdf")

                elif extension == "TXT":
                    textFile = open(
                        f"saves/tickets/{fullName}.txt", 'w', encoding='utf-8')
                    textFile.write(docx2txt.process(f"temp/{fullName}.docx"))

                os.remove(f"temp/{fullName}.docx")

            return True
        except Exception as e:
            print(e)
            return False

    def getCertificate(self, data, extension="DOCX"):
        try:
            document = Document("drafts/certificates.docx")
            table = document.tables[0]

            print("Hola")

            # Agrega el nombre completo
            fullName = f"{data[0]} {data[1]} {data[2]}"
            table.cell(1, 0).paragraphs[1].text = fullName

            # Agrega la categoría
            table.cell(2, 0).paragraphs[1].text = f"{data[8]}"

            if extension == "DOCX":
                document.save(f"saves/certificates/{fullName}.docx")

            else:
                document.save(f"temp/{fullName}.docx")

                if extension == "PDF":
                    convert(f"temp/{fullName}.docx",
                            f"saves/certificates/{fullName}.pdf")

                elif extension == "TXT":
                    textFile = open(
                        f"saves/certificates/{fullName}.txt", 'w', encoding='utf-8')
                    textFile.write(docx2txt.process(f"temp/{fullName}.docx"))

                os.remove(f"temp/{fullName}.docx")

            return True
        except Exception as e:
            print(e)
            return False

    def getAward(self, data, place, extension="DOCX"):
        try:
            document = Document("drafts/awards.docx")
            table = document.tables[0]

            # Agrega el nombre completo
            fullName = f"{data[0]} {data[1]} {data[2]}"
            table.cell(1, 0).paragraphs[1].text = fullName

            # Agrega la categoría
            table.cell(2, 0).paragraphs[1].text = f"{data[8]}"

            # Agrega el lugar del podio
            cell = table.cell(3, 0).paragraphs[1]

            if place == 1:
                cell.text = "1er lugar"
            elif place == 2:
                cell.text = "2do lugar"
            elif place == 3:
                cell.text = "3er lugar"

            if extension == "DOCX":
                document.save(f"saves/awards/{fullName}.docx")

            else:
                document.save(f"temp/{fullName}.docx")

                if extension == "PDF":
                    convert(f"temp/{fullName}.docx",
                            f"saves/awards/{fullName}.pdf")

                elif extension == "TXT":
                    textFile = open(
                        f"saves/awards/{fullName}.txt", 'w', encoding='utf-8')
                    textFile.write(docx2txt.process(f"temp/{fullName}.docx"))

                os.remove(f"temp/{fullName}.docx")

            return True
        except Exception as e:
            print(e)
            return False
